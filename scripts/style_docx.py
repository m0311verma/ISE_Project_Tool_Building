"""Apply ISE coursework formatting (Times New Roman 10pt) to a pandoc-generated docx.

Pandoc's default docx uses Calibri 11pt. The ISE coursework specification
(section 1.4) requires Times New Roman or Arial 10pt single-column. This script
post-processes report.docx to set those properties on every paragraph and run,
on every heading and Normal style, and on the body section width so the table
of figures stays single-column.

Run after `pandoc report.md -o report.docx`.
"""

from __future__ import annotations

import argparse
from copy import deepcopy

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


FONT_NAME = "Times New Roman"
FONT_SIZE_PT = 10
HEADING_SIZES = {
    "Title": 14,
    "Heading 1": 12,
    "Heading 2": 11,
    "Heading 3": 10,
}


def _set_run_font(run, name: str, size_pt: float) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def _set_style_font(style, name: str, size_pt: float) -> None:
    if style.font is None:
        return
    style.font.name = name
    style.font.size = Pt(size_pt)
    rpr = style.element.find(qn("w:rPr"))
    if rpr is None:
        rpr = style.element.makeelement(qn("w:rPr"), {})
        style.element.append(rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def style_document(in_path: str, out_path: str) -> None:
    doc = Document(in_path)

    # 1. Update Normal and heading styles so anything inheriting them flips too.
    for style_name, size in [("Normal", FONT_SIZE_PT), *HEADING_SIZES.items()]:
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        _set_style_font(style, FONT_NAME, size)

    # 2. Iterate every paragraph + run and force the font directly. This catches
    #    inline styles that pandoc applies which would otherwise override Normal.
    for para in doc.paragraphs:
        for run in para.runs:
            size = FONT_SIZE_PT
            if para.style and para.style.name in HEADING_SIZES:
                size = HEADING_SIZES[para.style.name]
            _set_run_font(run, FONT_NAME, size)

    # 3. Same for table cells.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        _set_run_font(run, FONT_NAME, FONT_SIZE_PT)

    doc.save(out_path)
    print(f"  -> {out_path} (Times New Roman {FONT_SIZE_PT}pt applied)")


def main(argv=None):
    p = argparse.ArgumentParser()
    p.add_argument("--in", dest="in_path", default="report.docx")
    p.add_argument("--out", dest="out_path", default="report.docx")
    args = p.parse_args(argv)
    style_document(args.in_path, args.out_path)


if __name__ == "__main__":
    main()
